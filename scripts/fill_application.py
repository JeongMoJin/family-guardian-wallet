"""
KFIP 2026 신청서 docx에 내용을 채워 넣는 스크립트.
양식(표 구조/스타일/폰트/병합/서명 영역) 100% 보존.
- 셀의 paragraph 텍스트만 in-place 갱신
- 다중 라인은 첫 paragraph 안에서 line break(<w:br/>) 로 처리

사용:
    python scripts/fill_application.py
    python scripts/fill_application.py --demo-url https://... --video-url https://...
"""

import argparse
import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches
from docx.table import _Cell


PROJECT = Path(__file__).resolve().parent.parent
SRC = PROJECT / "doc" / "Korea Financial Innovation Program 2026 신청서.docx"
DST = PROJECT / "doc" / "KFIP_2026_신청서_가디언월렛.docx"
DEFAULT_SIGNATURE = PROJECT / "doc" / "signature.png"


# ----- 입력 데이터 ----------------------------------------------------------

GITHUB_URL = "https://github.com/JeongMoJin/family-guardian-wallet"
SENIOR_ADDRESS = "rn3XdGVPxYboVZkdWzySGfi5K1LC2ym9gi"

PROJECT_INTRO = (
    "가디언월렛(Family Guardian Wallet)은 인지저하 시기 시니어의 자산을 가족 2/3 승인이 있어야만 인출되는 "
    "XRPL 네이티브 멀티시그 지갑이다. 시니어가 일정 금액 이상을 송금하려고 하면 등록된 가족 가디언 3명 중 "
    "2명이 모바일에서 승인해야만 트랜잭션이 testnet에 제출된다. 별도 컨트랙트 없이 XRPL의 SignerListSet과 "
    "multisign 으로 구현해 감사 리스크가 낮고, 3~5초 정산·낮은 수수료로 일상·소액 송금에도 적합하다. "
    "현재 testnet에서 SignerList 등록·1명 거부·2명 성공 흐름이 모두 검증되었고, 같은 흐름이 웹 데모에서 "
    "시니어/가디언 동시 화면으로 시연된다."
)

PROBLEM = (
    "한국은 2025년 초고령사회로 진입했고, 인지저하 초기와 후견 개시 사이의 ‘회색 지대’에 보이스피싱·자산 "
    "탈취·가족 분쟁 위험이 집중된다. 본인인증·OTP만 의존하는 현재 금융 시스템은 약해진 의사결정 능력을 "
    "보호하지 못하고, 성년후견 제도는 절차가 복잡하고 시점이 늦으며 1인 권한 집중의 위험이 있다. "
    "가디언월렛은 이 회색 지대에 가족 다수 합의를 트랜잭션 레이어 자체에 박는다. 큰돈이 나갈 때마다 가족이 "
    "한 번 더 보고 승인해야 비로소 자산이 움직이게 만들어, 단일 후견인의 권한 집중 없이도 시니어를 보호하는 "
    "구조다. 사용자 인사이트는 본 신청자가 운영 중인 시니어 인지건강 앱 ‘매일 두뇌건강’의 사용자군에서 직접 "
    "확보했다."
)

def prototype_lines(demo_url: str, video_url: str) -> list[str]:
    return [
        f"웹 데모: {demo_url}",
        f"데모 영상: {video_url}",
        f"GitHub: {GITHUB_URL}",
        f"시니어 testnet 주소: {SENIOR_ADDRESS}",
        "",
        "현재 동작 범위 (testnet 검증 완료):",
        "1) 시니어 1 + 가디언 3 testnet 계정 자동 생성, 시니어 계정에 SignerListSet (quorum=2, weight=1) 등록.",
        "2) CLI 데모 — 가디언 1명 서명: 거부 / 2명 서명: multisign 결합 후 testnet 제출 성공(tesSUCCESS).",
        "3) 웹 데모 — 시니어 화면(잔액·송금 폼) / 가디언 화면(대기 요청·승인 버튼) / 시연 모드(한 화면 분할).",
        "4) quorum 충족 시 서버가 자동으로 multisign 결합 + submitAndWait 으로 testnet 제출, explorer 링크가 시니어 화면에 노출.",
    ]


# ----- 셀 텍스트 안전 갱신 헬퍼 -------------------------------------------

def _replace_paragraph_text(paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell: _Cell, text_or_lines) -> None:
    """첫 paragraph 안에서만 작업해 셀 스타일을 보존한다.
    - 단일 문자열: 첫 paragraph 첫 run 에 그대로 입력.
    - 리스트(여러 줄): 첫 줄은 첫 run, 이후 줄은 line break 로 같은 paragraph 안에 추가.
    - 두 번째 이후 paragraph 텍스트는 비운다.
    """
    lines = text_or_lines if isinstance(text_or_lines, list) else [text_or_lines]
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        p = cell.add_paragraph()
        p.add_run(lines[0])
        for line in lines[1:]:
            r = p.add_run()
            r.add_break(WD_BREAK.LINE)
            r.add_text(line)
        return

    first = paragraphs[0]
    # 첫 줄 — 기존 run 의 첫 번째에 텍스트, 나머지 run 의 텍스트는 비움
    runs = first.runs
    if runs:
        runs[0].text = lines[0]
        for r in runs[1:]:
            r.text = ""
    else:
        first.add_run(lines[0])

    # 이후 줄 — line break + add_text 로 같은 paragraph 안에 이어 붙임
    for line in lines[1:]:
        r = first.add_run()
        r.add_break(WD_BREAK.LINE)
        r.add_text(line)

    # 첫 paragraph 외에 남아있는 paragraph 의 텍스트는 비운다 (구조는 유지)
    for p in paragraphs[1:]:
        for r in p.runs:
            r.text = ""


# ----- 본 작업 --------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--demo-url",
        default=os.environ.get("DEMO_URL")
        or "(배포 후 입력 — Render에서 발급되는 https://guardian-wallet-xxxx.onrender.com)",
    )
    parser.add_argument(
        "--video-url",
        default=os.environ.get("VIDEO_URL")
        or "(촬영·업로드 후 입력 — YouTube 링크 권장)",
    )
    parser.add_argument("--phone", default=os.environ.get("PHONE", ""))
    parser.add_argument(
        "--signature",
        default=str(DEFAULT_SIGNATURE) if DEFAULT_SIGNATURE.exists() else "",
        help="대표자 사인 PNG 경로 (기본: doc/signature.png 가 있으면 자동 사용)",
    )
    args = parser.parse_args()

    doc = Document(str(SRC))
    table = doc.tables[0]
    rows = table.rows

    # R2: 구분 — 개인 체크
    set_cell(rows[2].cells[4], "☑개인")

    # R3: 팀(기업)·개인명 / 사업분야
    set_cell(rows[3].cells[2], "진정모 (FLAG)")
    set_cell(rows[3].cells[7], "디지털 자산·시니어 인지건강 핀테크")

    # R4: 서비스
    set_cell(rows[4].cells[2], "가디언월렛 — 가족 2/3 승인 XRPL 멀티시그 지갑")

    # R5: 대표자 / 연락처 / 이메일
    set_cell(rows[5].cells[2], "진정모")
    set_cell(rows[5].cells[5], args.phone)
    set_cell(rows[5].cells[8], "wlswjdah123@naver.com")

    # R6: 담당자
    set_cell(rows[6].cells[2], "진정모")
    set_cell(rows[6].cells[5], args.phone)
    set_cell(rows[6].cells[8], "wlswjdah123@naver.com")

    # R7: 프로젝트 소개
    set_cell(rows[7].cells[2], PROJECT_INTRO)

    # R8: 해결하고자 하는 문제
    set_cell(rows[8].cells[2], PROBLEM)

    # R9: 깃허브
    set_cell(rows[9].cells[2], GITHUB_URL)

    # R10: XRPL 지갑 주소
    set_cell(rows[10].cells[2], f"{SENIOR_ADDRESS}  (시니어 계정 / XRPL testnet)")

    # R11: 프로토타입 (다중 라인)
    set_cell(rows[11].cells[2], prototype_lines(args.demo_url, args.video_url))

    # R13: 진행단계 — MVP개발 체크
    set_cell(rows[13].cells[2], "☑MVP개발")

    # R14: 서명 영역 — 날짜·대표자 이름·사인 이미지
    last_cell = rows[14].cells[0]
    for p in last_cell.paragraphs:
        t = p.text.strip()
        if t == "년 월 일":
            _replace_paragraph_text(p, "2026년 5월 12일")
        elif t.startswith("대표자"):
            if args.signature and Path(args.signature).exists():
                # "(인)" 텍스트 대신 사인 이미지를 같은 paragraph 끝에 inline 으로 삽입
                _replace_paragraph_text(p, "대표자: 진정모   ")
                run = p.add_run()
                run.add_picture(args.signature, width=Inches(0.55))
            else:
                _replace_paragraph_text(p, "대표자: 진정모   (인)")

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"saved: {DST}")


if __name__ == "__main__":
    main()
